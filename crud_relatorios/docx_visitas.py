import os
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from util.db import SQL


def gerar_relatorio_visitas():
    # conexão com o banco de dados
    sql = SQL(esquema='bd_gestao_visitantes')

    # consulta para "visitas por mês"
    cmd_visitas_mes = """
        SELECT YEAR(dta_visita) AS ano, MONTH(dta_visita) AS mes, COUNT(*) AS total_visitas
        FROM ta_visitas
        GROUP BY YEAR(dta_visita), MONTH(dta_visita)
        ORDER BY ano DESC, mes DESC;
    """
    visitas_mes = sql.get_list(cmd_visitas_mes)

    # consulta para obter o total de visitas
    cmd_visitas_total = """SELECT COUNT(*) AS total FROM ta_visitas GROUP BY YEAR(dta_visita) ORDER BY total;"""
    visitas_total = sql.get_int(cmd_visitas_total)

    # dicionário com os meses
    meses = {
        1: 'Janeiro', 2: 'Fevereiro', 3: 'Março', 4: 'Abril', 5: 'Maio', 6: 'Junho',
        7: 'Julho', 8: 'Agosto', 9: 'Setembro', 10: 'Outubro', 11: 'Novembro', 12: 'Dezembro'
    }

    doc = Document()

    logo_path = 'logo_CEUB_relatorios.png'
    try:
        doc.add_picture(logo_path, width=Inches(1.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    except FileNotFoundError:
        print(f"Arquivo de logo '{logo_path}' não encontrado. O relatório será gerado sem o logo.")

    # título do relatório
    title = doc.add_heading(level=0)
    wp = title.add_run("Relatório de Visitas")
    wp.font.color.rgb = RGBColor(67, 5, 78)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # para digitar texto no documento
    P1 = doc.add_paragraph("Este documento de relatório para as visitas do CEUB, com ele é possível visualizar as visitas por mês.")
    P1.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    P2 = doc.add_paragraph(f"De acordo com a extração de dados, se obteve um total de {visitas_total} visitas durante todo o ano de 2024.")
    P2.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    P3 = doc.add_paragraph("Para aumentar o número de visitas, é necessário o reforço de ações de marketing e divulgação, realização de eventos e promoções, além de melhorar a experiência do visitante com a disponibilidade de informações relevantes e segurança.")
    P3.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    P4 = doc.add_paragraph('Segue abaixo a tabela de visitas por mês, informando o fluxo de visitas por mês:')
    P4.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # tabela de visitas por mês
    t2 = doc.add_heading(level=1)
    wp2 = t2.add_run("Visitas por Mês")
    wp2.font.color.rgb = RGBColor(67, 5, 78)
    t2.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    table = doc.add_table(rows=1, cols=2)  # 1 linha de cabeçalho, 2 colunas (Mês e Visitas)
    table.style = 'Table Grid'

    # cabeçalho da tabela
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Mês'
    hdr_cells[1].text = 'Visitas'

    # cabeçalho
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)  # Branco
        tc_pr = cell._element.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), "6A0DAD")  # Roxo no formato hexadecimal
        tc_pr.append(shd)

    # total de visitas do ano
    total_visitas_ano = 0

    # criar e preencher a tabela
    for i in range(1, 13):  # Para cada mês de 1 a 12
        row_cells = table.add_row().cells
        row_cells[0].text = meses[i]  # Nome do mês

        # Verificar se há dados para o mês
        visitas_encontradas = next((row['total_visitas'] for row in visitas_mes if row['mes'] == i), None)
        if visitas_encontradas:
            row_cells[1].text = str(visitas_encontradas)
            total_visitas_ano += visitas_encontradas
        else:
            row_cells[1].text = 'SEM DADOS'
        for cell in row_cells:
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # linha do total de visitas
    total_row = table.add_row().cells
    total_row[0].text = 'Total de Visitas'
    total_row[1].text = str(total_visitas_ano)

    for cell in total_row:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        tc_pr = cell._element.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), "D8BFD8")  # Cor lilás clara no formato hexadecimal
        tc_pr.append(shd)

    # salvamento do arquivo
    relatorio_path = 'relatorio_visitas_2024.docx'
    doc.save(relatorio_path)
    print(f"Relatório salvo como '{relatorio_path}'.")

    # abre o arquivo automaticamente

    try:
        if os.name == 'nt':
            os.startfile(relatorio_path)
        elif os.name == 'posix':
            os.system(f'open {relatorio_path}')

    except Exception as e:
        print(f"Erro ao abrir o arquivo: {e}")